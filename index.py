import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import re
import argparse
import requests
from io import BytesIO
import urllib.parse
from urllib.parse import urlparse

def create_element(name):
    return OxmlElement(name)

def add_text_to_element(element, text):
    """Add text to element with proper encoding"""
    run = create_element('m:r')
    t = create_element('m:t')
    # Replace special characters and ensure proper spacing
    text = text.replace('Fixed', 'Fixed')  # Replace problematic word
    text = text.replace('-', '−')  # Use proper minus sign
    text = text.replace(' ', '\u0020')  # Ensure proper space character
    t.text = text
    run.append(t)
    element.append(run)

def create_fraction(num, den):
    """Create fraction element with proper text handling"""
    frac = create_element('m:f')
    num_elem = create_element('m:num')
    den_elem = create_element('m:den')
    
    # Add numerator text
    add_text_to_element(num_elem, num.strip('()'))
    
    # Add denominator text
    add_text_to_element(den_elem, den.strip('()'))
    
    frac.append(num_elem)
    frac.append(den_elem)
    return frac

def create_equation_element(formula):
    """Create Word equation element with proper text handling"""
    oMathPara = create_element('m:oMathPara')
    oMath = create_element('m:oMath')
    
    parts = formula.split('=')
    
    if len(parts) == 2:
        left_side, right_side = [p.strip() for p in parts]
        
        # Handle left side
        if '/' in left_side:
            num, den = left_side.split('/')
            oMath.append(create_fraction(num, den))
        else:
            add_text_to_element(oMath, left_side)
        
        # Add equals sign with proper spacing
        add_text_to_element(oMath, ' = ')
        
        # Handle right side
        if '/' in right_side:
            num, den = right_side.split('/')
            oMath.append(create_fraction(num, den))
        else:
            add_text_to_element(oMath, right_side)
    else:
        if '/' in formula:
            num, den = formula.split('/')
            oMath.append(create_fraction(num, den))
        else:
            add_text_to_element(oMath, formula)
    
    oMathPara.append(oMath)
    return oMathPara

def clean_formula(formula):
    """Convert LaTeX formula to plain math expression with proper character handling"""
    # Remove LaTeX wrappers
    formula = formula.replace('\\[', '').replace('\\]', '').strip()
    
    # Handle text blocks
    formula = re.sub(r'\\text{(.*?)}', lambda m: m.group(1).replace('-', '−'), formula)
    
    # Handle fractions
    while '\\frac{' in formula:
        match = re.search(r'\\frac{(.*?)}{(.*?)}', formula)
        if match:
            num = match.group(1)
            denom = match.group(2)
            formula = formula.replace(f'\\frac{{{num}}}{{{denom}}}', f'{num}/{denom}')
    
    # Handle other math symbols and ensure proper character encoding
    formula = formula.replace('\\times', '×')
    formula = formula.replace('-', '−')  # Use proper minus sign
    
    return formula.strip()

def download_image(url):
    """Download image from URL and return as bytes"""
    try:
        # Handle both HTTP and HTTPS URLs
        if not url.startswith(('http:', 'https:')):
            return None
        
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        return BytesIO(response.content)
    except Exception as e:
        print(f"Error downloading image from {url}: {str(e)}")
        return None

def add_image_to_doc(doc, image_url, max_width=6):
    """Add image to document with proper sizing"""
    image_stream = download_image(image_url)
    if image_stream:
        try:
            # Add the image with a maximum width of 6 inches
            doc.add_picture(image_stream, width=Inches(max_width))
            # Add a small space after the image
            doc.add_paragraph()
            return True
        except Exception as e:
            print(f"Error adding image to document: {str(e)}")
    return False

def convert_markdown_to_docx(input_file, output_file):
    # Read markdown content with explicit encoding
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Find all LaTeX formulas and store them
    formulas = {}
    formula_count = 0
    for match in re.finditer(r'\\\[(.*?)\\\]', md_content, re.DOTALL):
        placeholder = f'FORMULA_PLACEHOLDER_{formula_count}'
        formulas[placeholder] = match.group(0)
        md_content = md_content.replace(match.group(0), placeholder)
        formula_count += 1
    
    # Convert markdown to HTML
    html_content = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code']
    )
    
    # Parse HTML with explicit encoding
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create new Word document
    doc = Document()
    
    # Process each element
    for element in soup.find_all(['h3', 'p', 'table']):  # Removed 'img' from find_all
        if element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
            
        elif element.name == 'p':
            # Check if paragraph contains an image
            img = element.find('img')
            if img:
                # Add image if src attribute exists
                if img.get('src'):
                    add_image_to_doc(doc, img['src'])
            else:
                text = element.get_text()
                if 'FORMULA_PLACEHOLDER' in text:
                    for placeholder, original_formula in formulas.items():
                        if placeholder in text:
                            clean_text = clean_formula(original_formula)
                            p = doc.add_paragraph()
                            equation_element = create_equation_element(clean_text)
                            p._element.append(equation_element)
                            doc.add_paragraph()
                else:
                    doc.add_paragraph(text)
                
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                num_cols = len(rows[0].find_all(['td', 'th']))
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        for a in cell.find_all('a'):
                            a.replace_with(a.get_text())
                        table.cell(i, j).text = cell.get_text().strip()
            
            doc.add_paragraph()
    
    # Save the document with explicit encoding
    doc.save(output_file)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Convert markdown file to DOCX')
    parser.add_argument('-i', '--input', required=True, help='Input markdown file')
    parser.add_argument('-o', '--output', required=True, help='Output DOCX file')
    
    # Parse arguments
    args = parser.parse_args()
    
    try:
        convert_markdown_to_docx(args.input, args.output)
        print(f"Successfully converted {args.input} to {args.output}")
    except Exception as e:
        print(f"Error converting file: {str(e)}")